from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os
import tempfile

# Re-export screenshots via Pillow to strip all macOS metadata/DPI tags
_ORIG_SCREENSHOTS = "/data/screenshots"
_CLEAN_DIR = tempfile.mkdtemp()

for fname in sorted(os.listdir(_ORIG_SCREENSHOTS)):
    if fname.endswith('.png'):
        src = os.path.join(_ORIG_SCREENSHOTS, fname)
        dst = os.path.join(_CLEAN_DIR, fname)
        img = Image.open(src).convert('RGB')
        img.save(dst, format='PNG')
        print(f"Cleaned: {fname}")

SCREENSHOTS = _CLEAN_DIR
OUTPUT = "/data/Assignment4_Report.docx"

doc = Document()

# Set narrow margins to maximise image width
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

IMAGE_WIDTH = Inches(6)  # fill usable page width

def heading(text, level=1):
    doc.add_heading(text, level=level)

def body(text):
    doc.add_paragraph(text)

def code(text):
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_image(filename, caption=None):
    path = os.path.join(SCREENSHOTS, filename)
    doc.add_picture(path, width=IMAGE_WIDTH)
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(9)
        run.font.italic = True
    doc.add_paragraph()

# ── Title ──────────────────────────────────────────────────────────────────
title = doc.add_heading("Assignment 4 Report – CSD-3353 Web Applications in C#.NET", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Student: Qi Chen        Weight: 15%")
doc.add_paragraph("Topic: Final Enhancements – UI + Testing + Deployment")
doc.add_paragraph()

# ── CLOs ───────────────────────────────────────────────────────────────────
heading("Course Learning Outcomes")
body("CLO 1: Improve frontend styling using Bootstrap")
body("CLO 5: Debug and test application and deploy web application locally")

# ── 1. Bootstrap ───────────────────────────────────────────────────────────
heading("1. Bootstrap Styling")
body(
    "Bootstrap 5 was applied throughout the application to improve layout, "
    "readability, and visual clarity. Key changes include: a dark table header "
    "(table-dark), hover and striped rows, color-coded rating badges "
    "(green ≥8.0, yellow ≥6.0, red <6.0), genre displayed as badges, and "
    "Card-based layouts for Create, Edit, and Delete forms."
)

heading("1.1 Home Page", level=2)
add_image("01-home.png", "Figure 1 – Home Page")

heading("1.2 Movies List", level=2)
body("Responsive Bootstrap table with color-coded rating badges and styled action buttons.")
add_image("02-movies-list.png", "Figure 2 – Movies List")

heading("1.3 Create Movie Form", level=2)
body("Card component with a blue header and placeholder text for each input field.")
add_image("03-create-form.png", "Figure 3 – Create Form")

heading("1.4 Create Form – Filled In", level=2)
add_image("04-create-filled.png", "Figure 4 – Create Form (filled)")

heading("1.5 Movies List After Creating a Movie", level=2)
body("After successful submission the user is redirected to the updated movies list.")
add_image("05-after-create.png", "Figure 5 – After Create")

heading("1.6 Edit Movie Form", level=2)
body("Card component with a yellow header and pre-populated fields.")
add_image("06-edit-form.png", "Figure 6 – Edit Form")

heading("1.7 Delete Confirmation Page", level=2)
body("Danger-styled Card with a warning alert and full movie details before confirmation.")
add_image("07-delete-confirm.png", "Figure 7 – Delete Confirmation")

heading("1.8 Form Validation Errors", level=2)
body("Client-side and server-side validation errors shown inline using Bootstrap alert and text-danger spans.")
add_image("08-validation-errors.png", "Figure 8 – Validation Errors")

# ── 2. ILogger ─────────────────────────────────────────────────────────────
heading("2. Logging with ILogger")
body(
    "ILogger<MovieController> was injected into the MovieController constructor "
    "and used across all actions with three log levels:"
)
body("  • LogInformation – successful fetch, create, update, delete operations")
body("  • LogWarning      – null IDs, missing records, validation failures")
body("  • LogError        – database concurrency exceptions")

heading("Code sample", level=2)
code(
    'public MovieController(MovieDbContext context, ILogger<MovieController> logger)\n'
    '{\n'
    '    _context = context;\n'
    '    _logger  = logger;\n'
    '}\n\n'
    'public async Task<IActionResult> Index()\n'
    '{\n'
    '    _logger.LogInformation("Fetching all movies from database");\n'
    '    var movies = await _context.Movies.ToListAsync();\n'
    '    _logger.LogInformation("Retrieved {Count} movies", movies.Count);\n'
    '    return View(movies);\n'
    '}'
)

heading("Example console output", level=2)
code(
    '2026-04-08 17:09:36 info: MoviesManager.Controllers.MovieController\n'
    '      Fetching all movies from database\n'
    '2026-04-08 17:09:36 info: MoviesManager.Controllers.MovieController\n'
    '      Retrieved 3 movies\n'
    '2026-04-08 17:09:40 info: MoviesManager.Controllers.MovieController\n'
    '      Created new movie: Interstellar (Id=4)'
)

# ── 3. xUnit ───────────────────────────────────────────────────────────────
heading("3. xUnit Testing")
body(
    "A separate test project MoviesManager.Tests was created using xUnit, "
    "EF Core InMemory database, and Moq for logger mocking. "
    "18 tests are provided across two files."
)

heading("3.1 Model Validation Tests (MovieModelTests.cs) – 8 tests", level=2)
rows = [
    ("1", "Valid movie with all fields",           "Passes validation"),
    ("2", "Empty Title",                            "Fails – Title required"),
    ("3", "Rating above 10",                        "Fails – out of range"),
    ("4", "Year before 1888",                       "Fails – out of range"),
    ("5", "Title shorter than 2 characters",        "Fails – minimum length"),
    ("6", "Rating of exactly 0.0",                  "Passes validation"),
    ("7", "Rating of exactly 10.0",                 "Passes validation"),
    ("8", "Empty Director",                         "Fails – Director required"),
]
table = doc.add_table(rows=1 + len(rows), cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "#", "Test", "Expected Result"
for i, (num, test, result) in enumerate(rows, 1):
    row = table.rows[i].cells
    row[0].text, row[1].text, row[2].text = num, test, result
doc.add_paragraph()

heading("3.2 Controller Tests (MovieControllerTests.cs) – 10 tests", level=2)
rows2 = [
    ("1",  "Index() returns all movies",                        "View with 2 seeded movies"),
    ("2",  "Create() GET",                                      "Returns view"),
    ("3",  "Create() POST with valid data",                     "Saves to DB, redirects"),
    ("4",  "Edit() GET with valid id",                          "Returns view with movie"),
    ("5",  "Edit() GET with null id",                           "NotFound"),
    ("6",  "Edit() GET with non-existent id",                   "NotFound"),
    ("7",  "Delete() GET with valid id",                        "Returns view with movie"),
    ("8",  "DeleteConfirmed() POST",                            "Removes movie, redirects"),
    ("9",  "Create() POST with invalid model",                  "Returns view (no redirect)"),
    ("10", "Delete() GET with null id",                         "NotFound"),
]
table2 = doc.add_table(rows=1 + len(rows2), cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text, hdr2[1].text, hdr2[2].text = "#", "Test", "Expected Result"
for i, (num, test, result) in enumerate(rows2, 1):
    row = table2.rows[i].cells
    row[0].text, row[1].text, row[2].text = num, test, result
doc.add_paragraph()

heading("Run tests", level=2)
code("dotnet test MoviesManager.Tests/")

# ── 4. Deployment ──────────────────────────────────────────────────────────
heading("4. Deployment")

heading("4.1 dotnet publish", level=2)
code(
    "dotnet publish -c Release -o ./publish\n"
    "cd publish\n"
    "dotnet MoviesManager.dll"
)

heading("4.2 Docker", level=2)
code(
    "docker build -t movies-manager:a4 .\n"
    "docker run -d --name movies-app -p 8080:5000 \\\n"
    "  -e ASPNETCORE_URLS=http://+:5000 \\\n"
    "  movies-manager:a4"
)
body("Application accessible at: http://localhost:8080")

heading("4.3 Swagger API", level=2)
body("The REST API is documented and accessible at http://localhost:8080/swagger")
add_image("09-swagger.png", "Figure 9 – Swagger UI")

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
